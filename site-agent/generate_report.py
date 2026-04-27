#!/usr/bin/env python3
"""
CAYO Bar - Daily site-check Word report generator.

Reads results/latest.json, results/code_review.json (optional), and
results/api_probe.json (optional), and produces reports/cayo-site-report-<DATE>.docx
with full Hebrew RTL formatting using python-docx.
"""

from __future__ import annotations

import datetime as dt
import json
import os
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
LATEST = ROOT / "results" / "latest.json"
CODE_REVIEW_PATH = ROOT / "results" / "code_review.json"
API_PROBE_PATH = ROOT / "results" / "api_probe.json"
REPORTS_DIR = ROOT / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

if not LATEST.exists():
    raise SystemExit(f"ERROR: {LATEST} not found. Run agent.py first.")

DATA = json.loads(LATEST.read_text(encoding="utf-8"))
DATE = DATA.get("date") or dt.date.today().isoformat()

CODE_REVIEW = None
if CODE_REVIEW_PATH.exists():
    try:
        CODE_REVIEW = json.loads(CODE_REVIEW_PATH.read_text(encoding="utf-8"))
    except Exception:
        CODE_REVIEW = None

API_PROBE = None
if API_PROBE_PATH.exists():
    try:
        API_PROBE = json.loads(API_PROBE_PATH.read_text(encoding="utf-8"))
    except Exception:
        API_PROBE = None

# ---------------------------------------------------------------------------
# Palette
# ---------------------------------------------------------------------------

C = {
    "burgundy": RGBColor(0x4D, 0x14, 0x23),
    "teal": RGBColor(0x00, 0x85, 0x78),
    "cream": "F0E0C7",
    "orange": RGBColor(0xE3, 0x56, 0x32),
    "red": RGBColor(0xCB, 0x47, 0x47),
    "green": RGBColor(0x2E, 0x7D, 0x32),
    "amber": RGBColor(0xED, 0x6C, 0x02),
    "grey": RGBColor(0x88, 0x88, 0x88),
    "row_alt": "F7F2EA",
    "header_fill": "4D1423",
    "header_text": RGBColor(0xFF, 0xFF, 0xFF),
}

# ---------------------------------------------------------------------------
# XML / RTL helpers
# ---------------------------------------------------------------------------

def set_paragraph_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        pPr.append(bidi)


def set_run_rtl(run):
    rPr = run._r.get_or_add_rPr()
    if rPr.find(qn("w:rtl")) is None:
        rtl = OxmlElement("w:rtl")
        rtl.set(qn("w:val"), "1")
        rPr.append(rtl)


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def cell_borders(cell, color="CCCCCC", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_table_rtl(table):
    tblPr = table._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)


# ---------------------------------------------------------------------------
# Paragraph / run helpers
# ---------------------------------------------------------------------------

def add_par(doc, text="", *, bold=False, size=None, color=None,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=4, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    set_paragraph_rtl(p)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        if size: run.font.size = Pt(size)
        if bold: run.bold = True
        if color is not None: run.font.color.rgb = color
        set_run_rtl(run)
    return p


def add_runs(doc, runs, *, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             space_before=0, space_after=4, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    set_paragraph_rtl(p)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    for r in runs:
        run = p.add_run(r["text"])
        run.font.name = "Arial"
        run.font.size = Pt(r.get("size", 11))
        if r.get("bold"): run.bold = True
        if r.get("color") is not None: run.font.color.rgb = r["color"]
        set_run_rtl(run)
    return p


def add_section_heading(doc, text):
    return add_par(doc, text, bold=True, size=18, color=C["burgundy"],
                   space_before=16, space_after=8, style="Heading 1")


def add_sub_heading(doc, text):
    return add_par(doc, text, bold=True, size=14, color=C["teal"],
                   space_before=10, space_after=5, style="Heading 2")


def make_table(doc, col_widths):
    table = doc.add_table(rows=0, cols=len(col_widths))
    table.autofit = False
    table.allow_autofit = False
    set_table_rtl(table)
    return table


def set_cell_text(cell, text, *, bold=False, color=None, size=11,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT, fill=None):
    cell.text = ""
    if fill: shade_cell(cell, fill)
    cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = alignment
    set_paragraph_rtl(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text) if text is not None else "—")
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if bold: run.bold = True
    if color is not None: run.font.color.rgb = color
    set_run_rtl(run)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_row(table, cells_data):
    row = table.add_row()
    for i, data in enumerate(cells_data):
        c = row.cells[i]
        set_cell_text(c, data.get("text", ""),
                      bold=data.get("bold", False),
                      color=data.get("color"),
                      size=data.get("size", 11),
                      alignment=data.get("alignment", WD_ALIGN_PARAGRAPH.RIGHT),
                      fill=data.get("fill"))
    return row


def set_col_widths(table, inches):
    for row in table.rows:
        for i, c in enumerate(row.cells):
            c.width = Inches(inches[i])


# ---------------------------------------------------------------------------
# Status maps
# ---------------------------------------------------------------------------

def visual_status_he(status):
    return {
        "identical": "זהה לבייסליין",
        "baseline_created": "בייסליין נוצר",
        "minor_change": "שינוי קל",
        "noticeable_change": "שינוי מורגש",
        "major_change": "שינוי משמעותי",
        "skipped": "לא בוצע",
        None: "—",
    }.get(status, status or "—")


def visual_status_color(status):
    if status in ("identical", "baseline_created", "minor_change"):
        return C["green"]
    if status == "noticeable_change":
        return C["amber"]
    if status == "major_change":
        return C["red"]
    return C["grey"]


def status_color(status):
    if status == "ok": return C["green"]
    if status == "warn": return C["amber"]
    return C["red"]


def status_label(status):
    return {"ok": "תקין", "warn": "אזהרות", "fail": "כשלים"}.get(status, status)


def score_color(score):
    if score is None: return C["grey"]
    if score >= 90: return C["green"]
    if score >= 50: return C["amber"]
    return C["red"]


# ---------------------------------------------------------------------------
# Sections
# ---------------------------------------------------------------------------

def render_title(doc):
    add_par(doc, "CAYO Bar — דוח בדיקת תקינות יומי",
            bold=True, size=22, color=C["burgundy"], space_after=6)
    add_runs(doc, [
        {"text": "תאריך: ", "bold": True}, {"text": DATE},
        {"text": "  •  "},
        {"text": "כתובת: ", "bold": True}, {"text": DATA.get("url", "")},
    ])
    add_runs(doc, [
        {"text": "סטטוס כללי: ", "bold": True},
        {"text": status_label(DATA.get("overall", "—")),
         "bold": True, "color": status_color(DATA.get("overall", "ok"))},
    ])
    doc.add_paragraph()


def render_executive_summary(doc):
    s = DATA.get("summary") or {}
    overall = DATA.get("overall")
    if DATA.get("fatal_error"):
        text = f"שגיאה קריטית: לא ניתן היה לגשת לאתר. {DATA['fatal_error']}"
    elif overall == "ok":
        text = (f"✓ {s.get('functional_passed')}/{s.get('functional_total')} בדיקות פונקציונליות עברו. "
                f"ביצועים (mobile): {s.get('performance_score_mobile') or '—'}. "
                f"נגישות (mobile): {s.get('accessibility_score_mobile') or '—'}. "
                f"סטטוס ויזואלי: {visual_status_he(s.get('visual_status'))}. האתר תקין.")
    elif overall == "warn":
        text = (f"יש אזהרות שדורשות בדיקה. בדיקות פונקציונליות: {s.get('functional_passed')}/{s.get('functional_total')}. "
                f"ביצועים (mobile): {s.get('performance_score_mobile') or '—'}. "
                f"נגישות (mobile): {s.get('accessibility_score_mobile') or '—'}.")
    else:
        text = (f"נמצאו כשלים שדורשים תיקון. בדיקות פונקציונליות שעברו: {s.get('functional_passed')}/{s.get('functional_total')}.")

    # Append code review / API probe summary if present
    if CODE_REVIEW:
        cs = CODE_REVIEW.get("summary", {})
        sev = cs.get("by_severity", {})
        text += f" סקירת קוד: {cs.get('findings_total', 0)} ממצאים ({sev.get('high',0)} HIGH)."
    if API_PROBE:
        aps = API_PROBE.get("summary", {})
        text += f" API: {aps.get('high_severity', 0)} HIGH מתוך {aps.get('probed', 0)} מסלולים."

    add_section_heading(doc, "סיכום מנהלים")
    table = make_table(doc, [6.5])
    row = table.add_row()
    set_cell_text(row.cells[0], text, size=11, fill=C["cream"])
    row.cells[0].width = Inches(6.5)


def render_metrics(doc):
    s = DATA.get("summary") or {}
    perf = DATA.get("performance") or {}
    cwv_m = (perf.get("mobile") or {}).get("core_web_vitals") or {}
    cwv_d = (perf.get("desktop") or {}).get("core_web_vitals") or {}

    add_section_heading(doc, "מדדי מפתח")
    widths = [3.0, 1.75, 1.75]
    table = make_table(doc, widths)
    add_row(table, [
        {"text": "מדד", "bold": True, "color": C["header_text"], "fill": C["header_fill"]},
        {"text": "מובייל", "bold": True, "color": C["header_text"], "fill": C["header_fill"], "alignment": WD_ALIGN_PARAGRAPH.CENTER},
        {"text": "דסקטופ", "bold": True, "color": C["header_text"], "fill": C["header_fill"], "alignment": WD_ALIGN_PARAGRAPH.CENTER},
    ])

    def score_cell(val):
        if val is None:
            return {"text": "—", "alignment": WD_ALIGN_PARAGRAPH.CENTER, "color": C["grey"]}
        return {"text": str(val), "alignment": WD_ALIGN_PARAGRAPH.CENTER, "bold": True, "color": score_color(val)}

    rows_def = [
        ("ציון ביצועים (Lighthouse)", s.get("performance_score_mobile"), s.get("performance_score_desktop"), True),
        ("ציון נגישות (Lighthouse)", s.get("accessibility_score_mobile"), s.get("accessibility_score_desktop"), True),
        ("LCP — Largest Contentful Paint", cwv_m.get("lcp_display"), cwv_d.get("lcp_display"), False),
        ("FCP — First Contentful Paint", cwv_m.get("fcp_display"), cwv_d.get("fcp_display"), False),
        ("CLS — Cumulative Layout Shift", cwv_m.get("cls_display"), cwv_d.get("cls_display"), False),
        ("TBT — Total Blocking Time", cwv_m.get("tbt_display"), cwv_d.get("tbt_display"), False),
        ("TTI — Time to Interactive", cwv_m.get("tti_display"), cwv_d.get("tti_display"), False),
    ]
    for i, (label, vm, vd, is_score) in enumerate(rows_def):
        fill = C["row_alt"] if i % 2 else None
        if is_score:
            add_row(table, [
                {"text": label, "fill": fill, "bold": True},
                {**score_cell(vm), "fill": fill},
                {**score_cell(vd), "fill": fill},
            ])
        else:
            add_row(table, [
                {"text": label, "fill": fill},
                {"text": vm or "—", "fill": fill, "alignment": WD_ALIGN_PARAGRAPH.CENTER},
                {"text": vd or "—", "fill": fill, "alignment": WD_ALIGN_PARAGRAPH.CENTER},
            ])
    set_col_widths(table, widths)


def render_functional(doc):
    f = DATA.get("functional")
    add_section_heading(doc, "בדיקות פונקציונליות")
    if not f:
        add_par(doc, "בדיקות פונקציונליות לא רצו (שגיאה בגישה לאתר).", color=C["red"])
        return

    label_map = {
        "status_200": "החזר HTTP 200",
        "uses_https": "מוגש ב-HTTPS",
        "title_present": "תגית title קיימת",
        "title_contains_cayo": 'title מכיל "CAYO"',
        "lang_he": 'lang="he" על האלמנט html',
        "meta_description": "מטא-תיאור (meta description)",
        "og_title": "og:title להעדפת שיתוף",
        "og_image": "og:image להעדפת שיתוף",
        "h1_present": "כותרת H1 קיימת",
        "body_has_content": "גוף הדף מכיל תוכן (≥50 תווים)",
        "response_under_3s": "זמן תגובה < 3 שניות",
        "no_broken_images": "אין תמונות שבורות",
        "no_broken_links": "אין קישורים פנימיים שבורים",
    }
    widths = [2.4, 1.1, 3.0]
    table = make_table(doc, widths)
    add_row(table, [
        {"text": "בדיקה", "bold": True, "color": C["header_text"], "fill": C["header_fill"]},
        {"text": "תוצאה", "bold": True, "color": C["header_text"], "fill": C["header_fill"], "alignment": WD_ALIGN_PARAGRAPH.CENTER},
        {"text": "ערך בפועל", "bold": True, "color": C["header_text"], "fill": C["header_fill"]},
    ])
    for i, a in enumerate(f.get("assertions", [])):
        fill = C["row_alt"] if i % 2 else None
        passed = bool(a.get("passed"))
        add_row(table, [
            {"text": label_map.get(a["name"], a["name"]), "fill": fill},
            {"text": "✓ עבר" if passed else "✗ נכשל",
             "color": C["green"] if passed else C["red"],
             "bold": True, "alignment": WD_ALIGN_PARAGRAPH.CENTER, "fill": fill},
            {"text": str(a.get("actual", "") or "")[:120], "fill": fill},
        ])
    set_col_widths(table, widths)

    if f.get("broken_images"):
        add_sub_heading(doc, "תמונות שבורות")
        for b in f["broken_images"]:
            add_par(doc, f"• {b['url']} — {b.get('error') or 'HTTP ' + str(b.get('status'))}",
                    color=C["red"])
    if f.get("broken_links"):
        add_sub_heading(doc, "קישורים שבורים")
        for b in f["broken_links"]:
            add_par(doc, f"• {b['url']} — {b.get('error') or 'HTTP ' + str(b.get('status'))}",
                    color=C["red"])


def render_accessibility(doc):
    a11y = DATA.get("accessibility") or {}
    add_section_heading(doc, "נגישות (WCAG / Lighthouse)")
    score = (a11y.get("mobile") or {}).get("score")
    if score is None: score = (a11y.get("desktop") or {}).get("score")
    add_runs(doc, [
        {"text": "ציון נגישות: ", "bold": True},
        {"text": str(score if score is not None else "—"),
         "bold": True, "color": score_color(score)},
        {"text": " / 100"},
    ])
    failures = (a11y.get("mobile") or {}).get("failures") or (a11y.get("desktop") or {}).get("failures") or []
    if not failures:
        add_par(doc, "לא זוהו ממצאי נגישות שדורשים תשומת לב.", color=C["green"])
        return
    add_sub_heading(doc, f"ממצאים ({len(failures)})")
    widths = [2.2, 3.5, 0.8]
    table = make_table(doc, widths)
    add_row(table, [
        {"text": "מזהה Audit", "bold": True, "color": C["header_text"], "fill": C["header_fill"]},
        {"text": "בעיה", "bold": True, "color": C["header_text"], "fill": C["header_fill"]},
        {"text": "ציון", "bold": True, "color": C["header_text"], "fill": C["header_fill"], "alignment": WD_ALIGN_PARAGRAPH.CENTER},
    ])
    for i, fl in enumerate(failures[:30]):
        fill = C["row_alt"] if i % 2 else None
        sc = fl.get("score") or 0
        add_row(table, [
            {"text": fl.get("id", ""), "fill": fill, "size": 9},
            {"text": fl.get("title", ""), "fill": fill},
            {"text": f"{round(sc * 100)}%", "fill": fill,
             "alignment": WD_ALIGN_PARAGRAPH.CENTER, "bold": True,
             "color": C["red"] if sc < 0.5 else C["amber"]},
        ])
    set_col_widths(table, widths)


def render_performance(doc):
    perf = DATA.get("performance") or {}
    add_section_heading(doc, "ביצועים")
    add_par(doc,
            "הציונים נלקחים מ-Google PageSpeed Insights (Lighthouse).",
            color=C["grey"], size=10)
    failures = (perf.get("mobile") or {}).get("all_failures") or (perf.get("desktop") or {}).get("all_failures") or []
    if not failures:
        add_par(doc, "לא זוהו ממצאי שיפור משמעותיים.", color=C["green"])
        return
    add_sub_heading(doc, f"ממצאי שיפור ({len(failures)})")
    widths = [2.2, 3.4, 0.9]
    table = make_table(doc, widths)
    add_row(table, [
        {"text": "Audit", "bold": True, "color": C["header_text"], "fill": C["header_fill"]},
        {"text": "פירוט", "bold": True, "color": C["header_text"], "fill": C["header_fill"]},
        {"text": "ערך", "bold": True, "color": C["header_text"], "fill": C["header_fill"], "alignment": WD_ALIGN_PARAGRAPH.CENTER},
    ])
    for i, fl in enumerate(failures[:25]):
        fill = C["row_alt"] if i % 2 else None
        sc = fl.get("score") or 0
        display = fl.get("display") or f"{round(sc * 100)}%"
        add_row(table, [
            {"text": fl.get("id", ""), "fill": fill, "size": 9},
            {"text": fl.get("title", ""), "fill": fill},
            {"text": str(display)[:30], "fill": fill,
             "alignment": WD_ALIGN_PARAGRAPH.CENTER, "bold": True,
             "color": C["red"] if sc < 0.5 else C["amber"]},
        ])
    set_col_widths(table, widths)


def render_visual(doc):
    v = DATA.get("visual") or {}
    add_section_heading(doc, "בדיקה ויזואלית / רגרסיה")
    add_runs(doc, [
        {"text": "סטטוס: ", "bold": True},
        {"text": visual_status_he(v.get("status")),
         "bold": True, "color": visual_status_color(v.get("status"))},
    ])
    if v.get("today_hash"):
        add_runs(doc, [
            {"text": "Hash היום: ", "bold": True}, {"text": v.get("today_hash") or "—"},
            {"text": "  •  "},
            {"text": "Hash בייסליין: ", "bold": True}, {"text": v.get("baseline_hash") or "—"},
            {"text": "  •  "},
            {"text": "מרחק Hamming: ", "bold": True},
            {"text": f"{v.get('hamming_distance')}/64" if v.get("hamming_distance") is not None else "—"},
        ])
    today_path = v.get("today_path")
    if today_path and os.path.exists(today_path):
        add_sub_heading(doc, "צילום מסך — היום")
        try:
            doc.add_picture(today_path, width=Inches(3.6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            add_par(doc, f"לא ניתן לטעון צילום מסך: {e}", color=C["red"])


def render_code_review(doc):
    if not CODE_REVIEW:
        return
    doc.add_page_break()
    add_section_heading(doc, "סקירת קוד ולוגיקה")

    s = CODE_REVIEW.get("summary", {})
    sev = s.get("by_severity", {})
    add_runs(doc, [
        {"text": "ממצאים: ", "bold": True},
        {"text": str(s.get("findings_total", 0)), "bold": True, "color": C["burgundy"]},
        {"text": "  •  קבצים סרוקים: "},
        {"text": str(CODE_REVIEW.get("scanned", {}).get("files_total", 0)), "bold": True},
        {"text": "  •  חדש היום: "},
        {"text": str(s.get("new_today", 0)), "bold": True, "color": C["amber"]},
        {"text": "  •  נסגר אתמול: "},
        {"text": str(s.get("closed_since_yesterday", 0)), "bold": True, "color": C["green"]},
    ])
    add_runs(doc, [
        {"text": "לפי חומרה: ", "bold": True},
        {"text": "HIGH ", "color": C["red"], "bold": True}, {"text": str(sev.get("high", 0))},
        {"text": "  •  MEDIUM ", "color": C["amber"], "bold": True}, {"text": str(sev.get("medium", 0))},
        {"text": "  •  LOW ", "color": C["grey"], "bold": True}, {"text": str(sev.get("low", 0))},
    ])

    findings = CODE_REVIEW.get("findings", [])
    if not findings:
        add_par(doc, "לא נמצאו ממצאים — הקוד נקי.", color=C["green"])
        return

    for sev_key, sev_label, sev_color in (
        ("high", "חומרה גבוהה", C["red"]),
        ("medium", "חומרה בינונית", C["amber"]),
        ("low", "חומרה נמוכה", C["grey"]),
    ):
        bucket = [f for f in findings if f.get("severity") == sev_key]
        if not bucket:
            continue
        add_sub_heading(doc, f"{sev_label} ({len(bucket)})")

        widths = [1.0, 2.5, 0.6, 2.4]
        table = make_table(doc, widths)
        add_row(table, [
            {"text": "חומרה", "bold": True, "color": C["header_text"], "fill": C["header_fill"], "alignment": WD_ALIGN_PARAGRAPH.CENTER},
            {"text": "קובץ : שורה", "bold": True, "color": C["header_text"], "fill": C["header_fill"]},
            {"text": "כלל", "bold": True, "color": C["header_text"], "fill": C["header_fill"], "alignment": WD_ALIGN_PARAGRAPH.CENTER},
            {"text": "נושא", "bold": True, "color": C["header_text"], "fill": C["header_fill"]},
        ])
        for i, f in enumerate(bucket):
            fill = C["row_alt"] if i % 2 else None
            line_str = f"{f.get('file', '')}:{f.get('line_hint') or '?'}"
            sev_text = sev_key.upper() + (" *" if f.get("is_new") else "")
            add_row(table, [
                {"text": sev_text, "color": sev_color, "bold": True,
                 "alignment": WD_ALIGN_PARAGRAPH.CENTER, "fill": fill, "size": 9},
                {"text": line_str, "fill": fill, "size": 9},
                {"text": f.get("rule", ""), "fill": fill,
                 "alignment": WD_ALIGN_PARAGRAPH.CENTER, "bold": True, "size": 9},
                {"text": f.get("title", ""), "fill": fill, "size": 10},
            ])
        set_col_widths(table, widths)

        if sev_key in ("high", "medium"):
            limit = len(bucket) if sev_key == "high" else min(5, len(bucket))
            for f in bucket[:limit]:
                add_runs(doc, [
                    {"text": f"{f.get('file')}:{f.get('line_hint') or '?'}", "bold": True},
                    {"text": "  —  "},
                    {"text": f.get("title", ""), "bold": True, "color": sev_color},
                ], space_before=6)
                add_par(doc, f.get("detail", ""), size=10)
                if f.get("snippet"):
                    add_par(doc, f.get("snippet", "")[:240], size=9, color=C["grey"])


def render_api_probe(doc):
    if not API_PROBE:
        return
    add_section_heading(doc, "בדיקת API surface (ללא הרשאה)")
    s = API_PROBE.get("summary", {})
    add_runs(doc, [
        {"text": "מסלולים שנסרקו: ", "bold": True}, {"text": str(s.get("probed", 0))},
        {"text": "  •  HIGH: ", "bold": True, "color": C["red"]},
        {"text": str(s.get("high_severity", 0)), "bold": True, "color": C["red"]},
        {"text": "  •  MEDIUM: ", "bold": True, "color": C["amber"]},
        {"text": str(s.get("medium_severity", 0)), "bold": True, "color": C["amber"]},
    ])

    probes = API_PROBE.get("probes", [])
    if not probes:
        add_par(doc, "לא נסרקו מסלולים.", color=C["grey"])
        return

    widths = [1.7, 2.6, 0.8, 2.4]
    table = make_table(doc, widths)
    add_row(table, [
        {"text": "מסלול", "bold": True, "color": C["header_text"], "fill": C["header_fill"]},
        {"text": "תוצאה", "bold": True, "color": C["header_text"], "fill": C["header_fill"]},
        {"text": "סטטוס", "bold": True, "color": C["header_text"], "fill": C["header_fill"], "alignment": WD_ALIGN_PARAGRAPH.CENTER},
        {"text": "הערכה", "bold": True, "color": C["header_text"], "fill": C["header_fill"]},
    ])
    sev_color = {"high": C["red"], "medium": C["amber"], "low": C["grey"], "info": C["green"]}
    for i, p in enumerate(probes):
        fill = C["row_alt"] if i % 2 else None
        if p.get("skipped"):
            add_row(table, [
                {"text": p.get("path", ""), "fill": fill, "size": 9},
                {"text": "—", "fill": fill, "size": 9},
                {"text": "skip", "fill": fill, "alignment": WD_ALIGN_PARAGRAPH.CENTER, "size": 9},
                {"text": p.get("reason", ""), "fill": fill, "size": 9, "color": C["grey"]},
            ])
            continue
        resp = p.get("response", {})
        sev_v = p.get("severity", "info")
        add_row(table, [
            {"text": p.get("path", ""), "fill": fill, "size": 9},
            {"text": p.get("verdict", ""), "fill": fill, "size": 9,
             "color": sev_color.get(sev_v, C["grey"]), "bold": sev_v == "high"},
            {"text": str(resp.get("status", "—")), "fill": fill,
             "alignment": WD_ALIGN_PARAGRAPH.CENTER,
             "color": sev_color.get(sev_v, C["grey"]), "bold": True, "size": 10},
            {"text": p.get("reason", "")[:90], "fill": fill, "size": 9},
        ])
    set_col_widths(table, widths)


def render_appendix(doc):
    doc.add_page_break()
    add_section_heading(doc, "נספח — נתונים גולמיים")
    f = DATA.get("functional") or {}
    rows = [
        ("URL סופי", f.get("url_final")),
        ("סטטוס HTTP", f.get("status")),
        ("זמן תגובה", f"{f.get('elapsed_ms', 0)}ms"),
        ("Title", f.get("title")),
        ("Lang", f.get("lang")),
        ("Meta description", f.get("meta_description")),
        ("OG title", f.get("og_title")),
        ("OG image", f.get("og_image")),
        ("H1", " / ".join(f.get("h1") or []) or "—"),
        ("מספר קישורים", f.get("link_count")),
        ("מספר תמונות", f.get("image_count")),
    ]
    for label, val in rows:
        add_runs(doc, [{"text": f"{label}: ", "bold": True},
                       {"text": str(val if val is not None else '—')}])
    if f.get("body_text"):
        add_sub_heading(doc, "טקסט מתוך גוף הדף (חלקי)")
        add_par(doc, f["body_text"], size=10)


# ---------------------------------------------------------------------------
# Build & save
# ---------------------------------------------------------------------------

def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rfonts.set(qn("w:cs"), "Arial")

    render_title(doc)
    render_executive_summary(doc)
    render_metrics(doc)
    render_functional(doc)
    render_accessibility(doc)
    render_performance(doc)
    render_visual(doc)
    render_code_review(doc)
    render_api_probe(doc)
    render_appendix(doc)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ""
    run = fp.add_run(f"CAYO Site Agent  •  {DATE}")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = C["grey"]

    out_path = REPORTS_DIR / f"cayo-site-report-{DATE}.docx"
    doc.save(out_path)
    shutil.copy2(out_path, REPORTS_DIR / "cayo-site-report-latest.docx")
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
